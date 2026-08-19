import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def add_page_border(section, color="0F331F", sz="8"):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for b_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '20')
        b.set(qn('w:color'), color)
        pgBorders.append(b)
    sectPr.append(pgBorders)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(3, 105, 161)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(15, 118, 110)
    return p

def add_diagram_box(doc, title, flow_steps, bg_header="0F331F"):
    add_heading_styled(doc, f"❖ Workflow Architecture Diagram: {title}", level=2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r_t = p.add_run(f"ARCHITECTURAL FLOW: {title.upper()}\n")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(13)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(15, 51, 31)
    
    for idx, step in enumerate(flow_steps):
        p_step = cell.add_paragraph()
        p_step.paragraph_format.space_before = Pt(2)
        p_step.paragraph_format.space_after = Pt(4)
        
        step_prefix = f"► Step {idx+1}: "
        r_sp = p_step.add_run(step_prefix)
        r_sp.bold = True
        r_sp.font.name = 'Times New Roman'
        r_sp.font.size = Pt(11)
        r_sp.font.color.rgb = RGBColor(3, 105, 161)
        
        r_sb = p_step.add_run(step)
        r_sb.font.name = 'Times New Roman'
        r_sb.font.size = Pt(11)
        r_sb.font.color.rgb = RGBColor(30, 41, 59)
        
        if idx < len(flow_steps) - 1:
            p_arrow = cell.add_paragraph()
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arrow.paragraph_format.space_before = Pt(0)
            p_arrow.paragraph_format.space_after = Pt(0)
            r_arr = p_arrow.add_run("│\n▼")
            r_arr.bold = True
            r_arr.font.name = 'Times New Roman'
            r_arr.font.size = Pt(12)
            r_arr.font.color.rgb = RGBColor(15, 118, 110)

    set_table_borders(tbl, color="0F331F", sz="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_styled_table(doc, headers, rows_data, col_widths=None):
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Header row
    hdr_cells = tbl.rows[0].cells
    for idx, heading in enumerate(headers):
        hdr_cells[idx].text = heading
        set_cell_background(hdr_cells[idx], "0F331F")
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        if col_widths and idx < len(col_widths):
            hdr_cells[idx].width = Inches(col_widths[idx])
            
    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row_cells = tbl.rows[r_idx + 1].cells
        bg_color = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(30, 41, 59)
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = Inches(col_widths[c_idx])
                
    set_table_borders(tbl, color="CBD5E1", sz="4")
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(4)
    p_sp.paragraph_format.space_after = Pt(6)

def add_callout(doc, title, text, fill_hex="F0FDF4", border_hex="16A34A"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    p = cell.paragraphs[0]
    r_t = p.add_run(f"📌 {title}\n")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(11)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(15, 23, 42)
    
    r_b = p.add_run(text)
    r_b.font.name = 'Times New Roman'
    r_b.font.size = Pt(10)
    r_b.font.color.rgb = RGBColor(51, 65, 85)
    
    set_table_borders(tbl, color=border_hex, sz="6")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def create_portal_workflows_document():
    doc = Document()
    
    # Page setup - Margins & Page Border
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        add_page_border(section, color="0F331F", sz="8")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(6)
        footer_run = footer_p.add_run("© 2026 FIS Engineering Team — Sri Ramakrishna Engineering College, Coimbatore | SREC FIS V3.0")
        footer_run.font.name = 'Times New Roman'
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)

    # Official SREC Header Banner (same as PDF Reports)
    header_banner = os.path.abspath(os.path.join(os.path.dirname(__file__), '../client/public/srec-header-banner.png'))
    if os.path.exists(header_banner):
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_banner.paragraph_format.space_before = Pt(0)
        p_banner.paragraph_format.space_after = Pt(6)
        r_banner = p_banner.add_run()
        r_banner.add_picture(header_banner, width=Inches(7.0))

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sys.paragraph_format.space_before = Pt(0)
    p_sys.paragraph_format.space_after = Pt(8)
    r_sys = p_sys.add_run("FACULTY INFORMATION SYSTEM (SREC FIS V3.0)")
    r_sys.font.name = 'Times New Roman'
    r_sys.font.size = Pt(13)
    r_sys.bold = True
    r_sys.font.color.rgb = RGBColor(2, 132, 199)

    # Document Main Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(4)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("COMPLETE STRUCTURAL WORKFLOW SPECIFICATION FOR ALL 3 PORTALS")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(12)
    run_meta = meta_p.add_run(f"Document Generated: {datetime.datetime.now().strftime('%B %d, %Y')} | Version 3.0 | Status: Active Synchronized Workflow Specification")
    run_meta.font.name = 'Times New Roman'
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 116, 139)

    add_callout(doc, "AUTOMATIC LIVING DOCUMENTATION NOTICE", 
                "This single comprehensive Word document is automatically updated and regenerated whenever changes are made to user workflows, portal structures, approval lifecycles, or role permissions. The backend server automatically synchronizes this document upon system initialization and via automated CLI commands (`npm run update-docs`).", 
                fill_hex="EFF6FF", border_hex="0284C7")

    # Section 1
    add_heading_styled(doc, "1. System Architectural Overview & Authentication Framework", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("The SREC FIS V3.0 platform is engineered as a secure, high-performance, multi-tenant academic ERP designed to manage the entire lifecycle of faculty records, research outputs, institutional duties, accreditation metrics, and performance evaluations. The system operates across three distinct portals governed by a strict Role-Based Access Control (RBAC) matrix and centralized JSON Web Token (JWT) authentication.")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

    add_heading_styled(doc, "1.1 Role-Based Access Control (RBAC) Matrix", level=2)
    
    rbac_headers = ["User Role / Claim", "Portal Access", "Permission Scope", "Primary Functional Capabilities"]
    rbac_data = [
        ["faculty", "Portal 1: Regular Faculty", "Self Profile & Activities", "Profile logging, research/patent submissions, self FPI appraisal, AI CV builder, self dossier."],
        ["dept_admin (isHod: true)", "Portal 2: Dept Admin / HOD", "Department Wide", "Read-only faculty profile view, assign dept duties, evaluate FPI appraisals, bulk sign-off, dept dossiers."],
        ["admin (isInstitutionalAdmin)", "Portal 3: System Admin", "Institution Wide", "Full CRUD user management, rubric builder, final FPI approval, bulk sign-off, faculty transfers, accreditation suite."],
        ["principal / hr", "Portal 3: Executive Admin", "Institution Wide", "Executive appraisal reviews, institutional duty allocations, salary/experience certificate access, campus analytics."],
        ["isClubCoordinator: true", "Portal 1 + Club Module", "Assigned Club/Cell", "Log student events, guest speakers, budget utilization, participant rosters, club activity dossiers."],
        ["isSupervisorEligible (Ph.D)", "Portal 1 + Supervisor", "Research Guidance", "Register Ph.D. scholars guided, joint publications, thesis completion milestones, university allocations."]
    ]
    add_styled_table(doc, rbac_headers, rbac_data, [1.4, 1.6, 1.4, 2.6])

    add_diagram_box(doc, "Centralized JWT Authentication & Portal Access Routing", [
        "User enters credentials (Staff ID / Password / Role Selection) at /login.",
        "Backend verifies password hash via bcryptjs against staff_user or admin tables.",
        "Server generates 24-hour cryptographic JWT containing: staffId, role, isHod, isInstitutionalAdmin, isSupervisorEligible, isClubCoordinator.",
        "Client browser caches JWT in localStorage and attaches 'Authorization: Bearer <token>' to all subsequent API requests.",
        "React Router (App.jsx) and Role Guards dynamically mount Portal 1 (Faculty), Portal 2 (HOD), or Portal 3 (System Admin)."
    ])

    # Section 2
    add_heading_styled(doc, "2. Portal 1 — Regular Faculty Portal Structural Workflow", level=1)
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("Portal 1 is designed specifically for teaching and research faculty. It provides intuitive self-service workflows for managing professional credentials, scholarly publications, intellectual property, certifications, academic workloads, and annual self-performance evaluations.")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)

    fac_modules = [
        ("Workflow 1.1 — Authentication & Interactive Dashboard: ", 
         "Faculty authenticates and lands on Dashboard (/dashboard). The dashboard presents high-level summary cards (Total Publications, Grants Sanctioned, Certifications Completed, Pending Appraisals), quick-action links, and the Notification Bell for real-time announcements."),
        
        ("Workflow 1.2 — Personal Profile & Identification Data (/profile/personal): ", 
         "Faculty logs and updates personal demographic details including Date of Birth, Gender, Blood Group, Mobile, Emergency Contacts, Permanent/Communication Addresses, PAN, Aadhaar, APAAR ID, AICTE Faculty ID, and Anna University Faculty ID."),
        
        ("Workflow 1.3 — Academic Background & Research Identifiers (/profile/academic): ", 
         "Faculty views official academic metadata (Designation, Department, Date of Joining, SREC Experience, Total Experience). Faculty registers and maintains digital research identifiers: ORCID ID, Scopus Author ID, ResearcherID (Web of Science), Google Scholar Profile URL, and Vidwan Profile ID."),
        
        ("Workflow 1.4 — Document Vault & Official Identity File Uploads (/profile/documents): ", 
         "Faculty uploads scanned copies of essential identity documents (Aadhaar Card, PAN Card, Relieving Orders). Files are stored on the server under the secure organized hierarchy: /server/SREC/{Department}/{Staff_ID}/ with paths indexed in the database."),
        
        ("Workflow 1.5 — Educational Qualifications & Degree Certificates (/profile/education): ", 
         "Faculty submits chronologically ordered educational degrees (SSLC, HSC, Diploma, UG, PG, Ph.D.). Each qualification captures the degree name, specialization, institution, university, year of passing, percentage/CGPA, class/distinction, and scanned certificate proof."),
        
        ("Workflow 1.6 — Professional Society Memberships (/activities/memberships): ", 
         "Faculty registers active memberships in national and international professional bodies (e.g. IEEE, ACM, CSI, ISTE, IEI, IETE), logging membership ID, grade (Life Member, Senior Member, Fellow), and joining year."),
        
        ("Workflow 1.7 — Teaching Workload, Pass Percentages & Student Feedback: ", 
         "Faculty logs semester-wise course portfolios, subjects handled, student enrollment, laboratory courses, final semester examination pass percentages, and official student feedback scores (rated out of 5.0)."),
        
        ("Workflow 1.8 — Professional Development, Certifications & Events (/activities): ", 
         "Faculty records professional development milestones: (a) FDPs, STTPs, workshops, and seminars attended (/activities/interactions); (b) Online certifications from NPTEL, Swayam, Coursera, edX (/activities/certifications); (c) Invited guest lectures, session chairs, and resource person engagements (/activities/resource); (d) Honors, awards, and national fellowships (/activities/awards); (e) Workshops, conferences, and hackathons organized (/activities/events)."),
        
        ("Workflow 1.9 — R&D, Sponsored Grants, Seed Money, Patents & Publications: ", 
         "Faculty logs research achievements: (a) Sponsored research grants from DST, AICTE, DRDO, or industry partners (/activities/funding); (b) Institutional Seed Money grants (/activities/seed_money); (c) Patents, IPR filings, granted designs, and copyrights (/activities/ipr); (d) Journal publications (SCI, Scopus, UGC-CARE), conference proceedings, and book chapters (/activities/publications); (e) Authored/edited textbooks with ISBN (/activities/books); (f) Ph.D. scholars guided (/activities/scholars for scholars, /activities/supervisors for Ph.D. supervisors)."),
        
        ("Workflow 1.10 — Assigned Responsibilities & Institutional Service Viewer (/responsibilities): ", 
         "Faculty accesses the responsibilities module to review official departmental duties (e.g. Class Advisor, Lab In-charge, Timetable Coordinator) and institutional responsibilities (e.g. NAAC Coordinator, IQAC Member) assigned by HOD, Principal, HR, or System Admin."),
        
        ("Workflow 1.11 — Club Coordinator Sub-Portal Workflow (/activities/clubs): ", 
         "Faculty designated as Club Coordinators or Co-Coordinators access the Club Module to create club events, upload student participant rosters, log guest lectures, track budget utilization, and generate club annual activity dossiers."),
        
        ("Workflow 1.12 — Annual Faculty Performance Index (FPI) Self-Appraisal Submission (/appraisal): ", 
         "Faculty opens the Appraisal Form. The system auto-selects the completed academic year (2025-2026) and automatically calculates scores based on logged portal data: PART A (Teaching & Learning, max 60), PART B (Professional Development & Certifications, max 40), PART C (R&D, Grants & Publications, max 80), and PART D (Institutional & Departmental Responsibilities, max 20). Faculty verifies all auto-mapped proof documents via 👁️ View Proof buttons, sets future goals, applies self-score justifications if needed, and clicks 'Submit FPI Form' (Status becomes 'Submitted')."),
        
        ("Workflow 1.13 — 1-Click AI CV & Bio-Data Generator Engine (/cv-generator): ", 
         "Faculty opens the AI CV Generator to compile a comprehensive, industry-grade bio-data. Selects from 3 standardized institutional templates: (1) SREC Official Letterhead, (2) AICTE / Anna University Statutory Format, (3) Modern Technical Europass. Faculty selects AI Summary tone (Executive, Research-heavy, Teaching-focused), toggles section visibility, and exports vector-sharp PDF documents."),
        
        ("Workflow 1.14 — Self Academic Dossier & Multi-Format Report Generation (/reports): ", 
         "Faculty accesses Reports to generate aggregated dossiers, Excel matrices, and complete ZIP packages of all uploaded document evidence for promotions, accreditation, or annual reviews."),
        
        ("Workflow 1.15 — Settings & Automated White Background Profile Picture Standardization (/settings): ", 
         "Faculty uploads profile photos in Settings. The backend automatically processes images using AI portrait segmentation (rembg) to isolate the portrait and standardize the background to solid white (#FFFFFF), automatically updating the avatar across Navbar, Dashboard, and CV."),
        
        ("Workflow 1.16 — Real-Time Web Push Notification Subscription: ", 
         "Faculty clicks the Notification Bell in the Navbar to enroll the browser/mobile device for instant Web Push Notifications via Service Worker, receiving real-time alerts on appraisal status updates, HOD feedback, and administrative circulars."),

        ("Workflow 1.17 — Event Design & Certificate Generation Suite (/event-design-suite): ", 
         "Faculty organizing institutional events access the dedicated Event Design Suite to generate accredited Event Posters (P01-P05), Formal Invitations (I01-I05), and Bulk Participation Certificates (C01-C05). Features 1-click event loading from staff_event_organized, server-enforced department branding, server-locked 3-position Institutional Signatories (Faculty Coordinator, HOD, Principal), Excel/CSV participant spreadsheet import with real-time validation, live interactive HTML preview, vector PDF export, and JSZip bulk archive generation with deterministic numbering (SREC/<DEPT>/<YEAR>/<EVENT>/<SEQ>).")
    ]

    for t_wf, b_wf in fac_modules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run(t_wf)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(3, 105, 161)
        r2 = bp.add_run(b_wf)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)

    add_diagram_box(doc, "Faculty Annual FPI Appraisal Self-Submission Process", [
        "Faculty navigates to /appraisal for the active academic evaluation cycle (2025-2026).",
        "Appraisal Engine auto-aggregates all logged records (Pass %, Feedback, FDPs, Publications, Grants, Assigned Responsibilities).",
        "System computes auto-calculated marks across PART A (Max 60), PART B (Max 40), PART C (Max 80), PART D (Max 20).",
        "Faculty reviews claimed activities and verifies attached documents using 👁️ View Proof buttons.",
        "Faculty enters next-year professional goals, digitally signs the declaration, and clicks 'Submit FPI Form' (Status -> 'Submitted')."
    ])

    # Section 3
    add_heading_styled(doc, "3. Portal 2 — Department Admin & HOD Portal Structural Workflow", level=1)
    
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run("Portal 2 is tailored for Department Heads (HODs) and Department Administrators. It provides specialized tools for departmental governance, read-only faculty verification, departmental duty assignments, Part-by-Part appraisal scoring, batch sign-off, and aggregated departmental dossier reporting.")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(11)

    hod_modules = [
        ("Workflow 2.1 — Departmental Dashboard & Real-Time Analytics: ", 
         "HOD logs in to view department-specific analytics: total faculty strength, active research projects, total grant funds sanctioned, publication indexing breakdown (SCI vs. Scopus), and appraisal submission status progress bars."),
        
        ("Workflow 2.2 — Faculty Directory & Read-Only Profile Inspection (/admin/faculty): ", 
         "HOD views all department faculty members with access to 'View Dossier' and 'Docs' ZIP packages. To prevent unauthorized alterations, all form input fields in personal and academic profiles are strictly locked in read-only mode (disabled={auth.role === 'dept_admin'}). Sensitive HR documents (salary slips, relieving letters) are hidden from HOD view."),
        
        ("Workflow 2.3 — Departmental Additional Responsibilities Assignment (/responsibilities): ", 
         "HOD assigns department-level responsibilities (e.g. Class Advisor, Laboratory In-charge, Accreditation Criteria In-charge, Timetable In-charge, Placement Coordinator) to department faculty, specifying tenure and academic year."),
        
        ("Workflow 2.4 — Annual Performance Appraisal Review & Part-by-Part Evaluation (/appraisal): ", 
         "HOD monitors pending submissions (status === 'Submitted') from department faculty. HOD opens the Part-by-Part evaluation modal to inspect self-claimed scores against actual uploaded proofs, enters HOD evaluated scores for PART A, PART B, PART C, and PART D, provides qualitative comments, and clicks 'Approve & Forward to Principal/HR'."),
        
        ("Workflow 2.5 — Auto-Mapped Activity Verification & Evidence Inspection: ", 
         "Within the appraisal interface, HOD inspects original uploaded proofs (journal PDFs, grant sanction orders, patent certificates, FDP certificates) by clicking 👁️ View Proof next to each mapped claim, guaranteeing strict audit compliance."),
        
        ("Workflow 2.6 — Batch Processing & Bulk Sign/Approval Workflow (✍️ Sign & Approve Selected): ", 
         "For efficient department-wide processing, HOD selects multiple or all pending forms via checkboxes and clicks '✍️ Sign & Approve Selected'. The system opens the Digital Signature Modal displaying signer credentials (HOD name, ISO timestamp, client IP address), applies bulk remarks, and submits in a single batch (/api/faculty/appraisals/bulk-hod-sign-approve). Status transitions to 'HOD Approved' and notification emails/pushes are sent to faculty."),
        
        ("Workflow 2.7 — Form Rejection & Revision Return Loop (Returned for Correction): ", 
         "If discrepancies or missing proof files are detected, HOD can return the appraisal to the faculty member with specific revision notes. Form status updates to 'Returned for Correction', unlocking the form for faculty edits."),
        
        ("Workflow 2.8 — Department-Wide Dossier & Accreditation Report Generation (/reports): ", 
         "HOD generates consolidated department reports: NBA / NAAC criteria matrices, annual publication dossiers, funded project portfolios, and departmental appraisal summary sheets in PDF and Excel formats.")
    ]

    for t_wf, b_wf in hod_modules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run(t_wf)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(15, 51, 31)
        r2 = bp.add_run(b_wf)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)

    add_diagram_box(doc, "HOD Part-by-Part Appraisal Verification & Bulk Score Evaluation", [
        "HOD monitors department submissions queue at /appraisal (Status: 'Submitted').",
        "HOD inspects attached evidence for all claimed activities via 👁️ View Proof.",
        "HOD executes individual Part-by-Part scoring or initiates batch processing via '✍️ Sign & Approve Selected'.",
        "System records official digital signature (HOD name, timestamp, client IP) and applies evaluation remarks.",
        "Appraisal status transitions to 'HOD Approved' and forwards the form to the Principal/Executive queue."
    ])

    # Section 4
    add_heading_styled(doc, "4. Portal 3 — System Admin & Executive Portal Structural Workflow", level=1)
    
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    r4 = p4.add_run("Portal 3 equips the Principal, Human Resources, and System Administrators with institution-wide control, full CRUD user management, dynamic form construction, appraisal rubric configurability, multi-mode bulk sign-offs, and automated faculty department transfer execution.")
    r4.font.name = 'Times New Roman'
    r4.font.size = Pt(11)

    adm_modules = [
        ("Workflow 3.1 — Executive Institutional Dashboard & Accreditation Suite (/accreditation-suite): ", 
         "Executive Admins access institution-wide analytics and real-time NBA / NAAC compliance scorecards. Visualizes department-by-department comparisons, grant inflows, Ph.D. completions, and total institutional publications."),
        
        ("Workflow 3.2 — Institution-Wide User Management & Full Administrative CRUD: ", 
         "Admin manages all user accounts across System Admins (/admin/system-admins), Dept Admins (/admin/dept-admins), Faculty Directory (/admin/faculty), and Club Coordinators (/admin/clubs). Admins have complete authority to create accounts, modify personal/academic details, update designations, reset passwords, and mark faculty as relieved (is_relieved === 1)."),
        
        ("Workflow 3.3 — Institution-Level Responsibilities Management (/responsibilities): ", 
         "Principal/HR/Admin assigns institutional responsibilities (e.g. NAAC Steering Committee Coordinator, IQAC Director, Anti-Ragging Committee Head, Chief Warden, Controller of Examinations) to faculty across any department."),
        
        ("Workflow 3.4 — Dynamic Page & Custom Form Builder Engine (/admin/dynamic-pages): ", 
         "Admin creates custom survey forms, feedback questionnaires, or data collection pages. Admins configure form fields, validation rules, target department visibility, and role permissions dynamically without code changes."),
        
        ("Workflow 3.5 — Appraisal Form Builder & Rubric Configurator (/appraisal -> Form Builder): ", 
         "Admin configures master FPI evaluation rubrics: (a) Designation Overrides: Customizes unit marks and maximum caps for Assistant Professor, Associate Professor, Professor, and Professor & Head; (b) Threshold Brackets: Sets cutoff criteria for student feedback (e.g. 4.0 cutoff), pass % brackets (e.g. 80% cutoff), and publication splits; (c) Custom PART Additions: Dynamically adds new evaluation sections (e.g. PART_E, PART_F) with custom titles and mark ceilings."),
        
        ("Workflow 3.6 — Final Executive Appraisal Approval & Institutional Bulk Sign-Off: ", 
         "Executive Admins review HOD-approved submissions (status === 'HOD Approved') or submitted forms across all departments. Admins inspect proofs (👁️ View Proof), assign final executive scores, enter increment/promotion remarks, and finalize individual records. For bulk actions, Admin opens the Multi-Mode Bulk Sign & Approve Modal (Auto-detect, Final Executive Approval, or HOD Stage bypass), applies digital signatures, and finalizes all selected records in one click (Status -> 'Final Approved')."),
        
        ("Workflow 3.7 — Faculty Department Transfer Engine (/api/admin/faculty/transfer): ", 
         "Admin executes atomic faculty transfers between departments. The backend records immutable audit logs in staff_department_history, updates staff_academics, moves physical storage directory from /SREC/{old_dept}/{staff_id}/ to /SREC/{new_dept}/{staff_id}/, remaps DB document paths, and re-queues the faculty member under the new department HOD."),
        
        ("Workflow 3.8 — System Maintenance, DB Backups & Institution Bulk Export: ", 
         "Admin generates institution-wide bulk dossier ZIP packages containing all faculty PDF reports and document proofs, and executes database maintenance operations.")
    ]

    for t_wf, b_wf in adm_modules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run(t_wf)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(185, 28, 28)
        r2 = bp.add_run(b_wf)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)

    add_diagram_box(doc, "Faculty Department Transfer & Directory Relocation Lifecycle", [
        "System Admin initiates faculty transfer at /api/admin/faculty/transfer (Source Dept -> Target Dept).",
        "Backend creates immutable audit record in staff_department_history with effective date.",
        "Backend updates Department field in staff_academics and syncs admin_dep table if user held admin roles.",
        "Physical Disk Mover (moveFacultyDirectory) relocates files: /SREC/OldDept/ID/ ➔ /SREC/NewDept/ID/.",
        "Database Engine updates path strings in all activity tables and re-queues faculty under New Department HOD."
    ])

    # Section 5
    add_heading_styled(doc, "5. Cross-Portal Unified Lifecycles & Interaction Maps", level=1)
    
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(6)
    r5 = p5.add_run("The three portals operate in seamless synchronization. Critical institutional processes flow seamlessly across portals through well-defined multi-stage state machines.")
    r5.font.name = 'Times New Roman'
    r5.font.size = Pt(11)

    add_heading_styled(doc, "5.1 End-to-End Annual Performance Appraisal (FPI) Lifecycle", level=2)
    add_diagram_box(doc, "Three-Tier Multi-Stage Appraisal Approval Workflow", [
        "Stage 1 (Portal 1 - Faculty): Faculty fills self-scores across PART A, B, C, D, attaches proofs, and submits form (Status: 'Submitted').",
        "Stage 2 (Portal 2 - HOD): HOD reviews auto-mapped proofs, inputs evaluated marks, or executes bulk sign-off (Status: 'HOD Approved' or 'Returned for Correction').",
        "Stage 3 (Portal 3 - Executive Admin / Principal): Executive Admin conducts final institutional review, assigns executive scores, and clicks 'Final Approve' (Status: 'Final Approved').",
        "Stage 4 (Automated Dispatch): System locks form against further editing, generates final composite evaluation report, and dispatches Web Push & Email notifications to faculty."
    ])

    add_heading_styled(doc, "5.2 End-to-End Institutional Accreditation & Bulk Dossier Generation", level=2)
    add_diagram_box(doc, "Accreditation Data Harvesting & Dossier Archive Lifecycle", [
        "Stage 1 (Portal 1 - Faculty Data Entry): Faculty continuously logs publications, grants, patents, FDPs, and uploads certificates.",
        "Stage 2 (Portal 2 - Department Verification): HOD verifies evidence authenticity and monitors departmental NBA/NAAC criteria completion.",
        "Stage 3 (Portal 3 - Accreditation Suite): Institutional IQAC/NAAC coordinators access /accreditation-suite to view real-time compliance matrices.",
        "Stage 4 (Bulk Zip Exporter): Exporter engine (/api/reports/bulk-dossier-zip) streams compressed ZIP containing structured folders and vector PDFs."
    ])

    # Section 6
    add_heading_styled(doc, "6. Comprehensive Portal Feature & Permission Comparison Matrix", level=1)
    
    matrix_headers = ["Functional Feature / Module", "Portal 1 (Faculty)", "Portal 2 (HOD / Dept Admin)", "Portal 3 (System Admin / Principal)"]
    matrix_data = [
        ["Authentication & Login", "Staff ID + Password", "Staff ID + Password (isHod: true)", "Admin Username + Password"],
        ["Personal & Academic Details", "Edit & Save Self Profile", "Read-Only Inspection (Disabled)", "Full Administrative CRUD"],
        ["Educational Qualifications", "Add/Edit Self Qualifications", "Read-Only Inspection", "Full Administrative CRUD"],
        ["Document Vault Uploads", "Upload Self Documents", "Read-Only (HR Docs Restricted)", "Full Access + HR Dossier View"],
        ["Research, Patents & Grants", "Submit Self R&D Records", "View Dept Records & Proofs", "Full CRUD + Institutional Summary"],
        ["Club Coordinator Module", "Assigned Clubs Only", "Read-Only Dept View", "Create Clubs & Assign Coordinators"],
        ["Responsibilities Assignment", "View Assigned Duties", "Assign Departmental Duties", "Assign Institutional Duties"],
        ["Appraisal Self-Submission", "Self-Score & Submit (Part A-D)", "Cannot Self-Submit in Admin View", "Cannot Self-Submit in Admin View"],
        ["Appraisal Evaluation / Scoring", "None (Self-Evaluation Only)", "Part-by-Part HOD Scoring", "Final Executive Scoring & Overrides"],
        ["Appraisal Bulk Sign & Approve", "Not Available", "✍️ Bulk HOD Sign & Approve", "✍️ Multi-Mode Bulk Executive Sign"],
        ["Appraisal Return for Correction", "Receive Revision Request", "Return to Faculty with Remarks", "Return to HOD or Faculty"],
        ["Rubric & Form Builder", "Not Available", "Not Available", "⚙️ Config Rubric, Brackets & Parts"],
        ["AI CV & Bio-Data Generator", "Generate 3 Formats + AI Summary", "View & Download Faculty CVs", "View & Download Faculty CVs"],
        ["Dossier & Report Generation", "Self Dossier (PDF/Excel/ZIP)", "Dept Aggregated Dossiers", "Institution Bulk Dossier ZIP Exporter"],
        ["Faculty Department Transfer", "Not Available", "Not Available", "Atomic Transfer + Disk & DB Remap"],
        ["Accreditation Suite (NBA/NAAC)", "View Personal Contribution", "View Department Compliance", "Institution Executive Scorecards"]
    ]
    add_styled_table(doc, matrix_headers, matrix_data, [2.0, 1.6, 1.7, 1.7])

    # Section 7
    add_heading_styled(doc, "7. Automated Future Synchronization & Maintenance System", level=1)
    
    p7 = doc.add_paragraph()
    p7.paragraph_format.space_after = Pt(6)
    r7 = p7.add_run("To ensure that this single comprehensive Word document remains permanently synchronized with all future codebase evolutions, the repository implements an automated multi-layered update infrastructure:")
    r7.font.name = 'Times New Roman'
    r7.font.size = Pt(11)

    sync_points = [
        ("1. Dedicated Generator Script: ", "The python script scripts/generate_portal_workflows_doc.py is the single source of truth for generating docs/Complete_3_Portals_Workflow_Guide.docx, utilizing python-docx with official SREC branding, color palettes, and structured tables."),
        ("2. Backend Server Startup Auto-Sync: ", "In server/db.js and server/server.js, the backend automatically invokes generate_portal_workflows_doc.py during database initialization on server launch, ensuring that local and production deployments maintain synchronized documentation."),
        ("3. Unified NPM Documentation Commands: ", "Developers can regenerate all system documentation simultaneously by running `npm run update-docs` from the server directory, which sequentially regenerates the Workflow Guide, Database Schema, System Constraints, and Technical File Guide."),
        ("4. Workspace AI Agent Rules (.agents/AGENTS.md): ", "Mandatory workspace agent rules enforce that any change to role permissions, approval lifecycles, or menu structures must trigger an automated execution of the generator script before committing changes.")
    ]

    for t_s, b_s in sync_points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(t_s)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(15, 23, 42)
        r2 = bp.add_run(b_s)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10)

    # Save document
    docs_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../docs'))
    os.makedirs(docs_dir, exist_ok=True)
    out_path = os.path.join(docs_dir, "Complete_3_Portals_Workflow_Guide.docx")
    doc.save(out_path)
    print(f"Complete 3-Portal Workflow Word document successfully generated at: {out_path}")

if __name__ == "__main__":
    create_portal_workflows_document()
