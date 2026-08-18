import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def add_page_border(section, color="0F331F", sz="8"):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for b_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '20')
        b.set(qn('w:color'), color)
        pgBorders.append(b)
    sectPr.append(pgBorders)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(3, 105, 161)
    return p

def add_diagram_box(doc, title, flow_steps):
    add_heading_styled(doc, f"❖ System Architecture Diagram: {title}", level=2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r_t = p.add_run(f"SYSTEM RULE DIAGRAM: {title.upper()}\n")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(14)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(15, 51, 31)
    
    for idx, step in enumerate(flow_steps):
        p_step = cell.add_paragraph()
        p_step.paragraph_format.space_before = Pt(2)
        p_step.paragraph_format.space_after = Pt(4)
        
        step_prefix = f"► Stage {idx+1}: "
        r_sp = p_step.add_run(step_prefix)
        r_sp.bold = True
        r_sp.font.name = 'Times New Roman'
        r_sp.font.size = Pt(14)
        r_sp.font.color.rgb = RGBColor(3, 105, 161)
        
        r_sb = p_step.add_run(step)
        r_sb.font.name = 'Times New Roman'
        r_sb.font.size = Pt(14)
        r_sb.font.color.rgb = RGBColor(30, 41, 59)
        
        if idx < len(flow_steps) - 1:
            p_arrow = cell.add_paragraph()
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arrow.paragraph_format.space_before = Pt(0)
            p_arrow.paragraph_format.space_after = Pt(0)
            r_arr = p_arrow.add_run("│\n▼")
            r_arr.bold = True
            r_arr.font.name = 'Times New Roman'
            r_arr.font.size = Pt(14)
            r_arr.font.color.rgb = RGBColor(15, 118, 110)

    set_table_borders(tbl, color="0F331F", sz="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def create_system_constraints_document():
    doc = Document()
    
    # Page setup - Margins & Page Border
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        add_page_border(section, color="0F331F", sz="8")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(6)
        footer_run = footer_p.add_run("© 2026 FIS Team - Sri Ramakrishna Engineering College, Coimbatore | SREC FIS V3.0")
        footer_run.font.name = 'Times New Roman'
        footer_run.font.size = Pt(11)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)

    # Official SREC Header Banner (same as PDF Reports)
    header_banner = os.path.abspath(os.path.join(os.path.dirname(__file__), '../client/public/srec-header-banner.png'))
    if os.path.exists(header_banner):
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_banner.paragraph_format.space_before = Pt(0)
        p_banner.paragraph_format.space_after = Pt(6)
        r_banner = p_banner.add_run()
        r_banner.add_picture(header_banner, width=Inches(7.0))

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sys.paragraph_format.space_before = Pt(0)
    p_sys.paragraph_format.space_after = Pt(10)
    r_sys = p_sys.add_run("FACULTY INFORMATION SYSTEM (SREC FIS V3.0)")
    r_sys.font.name = 'Times New Roman'
    r_sys.font.size = Pt(14)
    r_sys.bold = True
    r_sys.font.color.rgb = RGBColor(2, 132, 199)

    # Document Main Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("SYSTEM CONSTRAINTS, RULES & PORTAL FUNCTIONALITY MANUAL")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(14)
    run_meta = meta_p.add_run(f"Document Generated: {datetime.datetime.now().strftime('%B %d, %Y')} | Version 3.0 | Status: Active System Rules Specification")
    run_meta.font.name = 'Times New Roman'
    run_meta.font.size = Pt(11)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Section 1
    add_heading_styled(doc, "1. Executive Summary & Portal Classification", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("The Sri Ramakrishna Engineering College Faculty Information System (SREC FIS V3.0) operates across three distinct portal environments, governed by strict authorization constraints, menu visibility matrix, role elevation logic, and workflow governance.")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    table1 = doc.add_table(rows=1, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False
    
    hdr_cells = table1.rows[0].cells
    headers = ["Portal Environment", "Target User Roles", "Operational Scope & Access Rules"]
    widths = [Inches(1.8), Inches(1.8), Inches(3.4)]
    
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(14)
        set_cell_background(hdr_cells[i], "0F331F")
        set_cell_margins(hdr_cells[i])
        hdr_cells[i].width = widths[i]

    portals_data = [
        ("1. Faculty Portal", "role: 'faculty'", "Self-service profile logging, personal activities, R&D submissions, assigned responsibilities view, self FPI appraisal submission."),
        ("2. Dept Admin / HOD Portal", "role: 'dept_admin'\nisHod: true", "Departmental faculty oversight, read-only faculty verification, departmental additional responsibilities assignment, HOD Part-by-Part appraisal review & score evaluation."),
        ("3. Executive & System Admin Portal", "role: 'admin'\nrole: 'principal'\nrole: 'hr'", "Institution-wide administration, full CRUD across all faculty records, Dynamic Form Builder & Rubrics Configurator, designation scoring parameter overrides, final executive appraisal approval.")
    ]

    for row_idx, (p_name, p_roles, p_scope) in enumerate(portals_data):
        row_cells = table1.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate([p_name, p_roles, p_scope]):
            row_cells[i].text = val
            p_cell = row_cells[i].paragraphs[0]
            p_cell.runs[0].font.name = 'Times New Roman'
            p_cell.runs[0].font.size = Pt(14)
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i])
            row_cells[i].width = widths[i]

    set_table_borders(table1)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 2
    add_heading_styled(doc, "2. User Roles, Elevation Flags & Automatic Designation Rules", level=1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("The system dynamically evaluates user credentials and designation strings to elevate permissions and toggle menu items:")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    rules = [
        ("Doctorate / Ph.D Designation Rule: ", "If a faculty member's name contains 'Dr.' or 'Dr ', or Ph.D qualification is recorded, the Research Supervisor menu item (/activities/supervisors) is automatically enabled in place of the Research Scholar menu item (/activities/scholars)."),
        ("HOD Designation Rule: ", "If a faculty member's designation contains 'Head' or 'HOD', the isHod flag is set to true, granting Departmental Responsibility Assignment and HOD Appraisal Review privileges."),
        ("Executive Admin Rule: ", "If a faculty member's designation contains 'Principal' or 'HR', the isInstitutionalAdmin flag is set to true, granting Institutional Responsibility Assignment, Executive Appraisal Review, and Form Builder access."),
        ("Club Coordinator Rule: ", "If a faculty member is assigned as an active coordinator or co-coordinator in staff_club or clubs, the isClubCoordinator flag is set to true, enabling the Clubs Activity Organized menu item."),
        ("Relieved Faculty Constraint: ", "If is_relieved === 1 in staff_user or staff_personal, login authentication is strictly blocked, revoking all system access.")
    ]

    for title_text, body_text in rules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(title_text)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(15, 118, 110)
        r2 = bp.add_run(body_text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(14)

    add_diagram_box(doc, "Dynamic Permission Elevation & Menu Toggle Rule Logic", [
        "Login credentials verified against MySQL database.",
        "System parses Designation string: Contains 'Head'/'HOD' ➔ Elevate isHod = true.",
        "System parses Designation string: Contains 'Principal'/'HR' ➔ Elevate isInstitutionalAdmin = true.",
        "System checks Qualification/Name: Contains 'Dr.' or Ph.D ➔ Enable Research Supervisor menu.",
        "System checks Club Assignments: Assigned Coordinator/Co-Coordinator ➔ Enable Clubs Activity menu."
    ])

    # Section 3
    add_heading_styled(doc, "3. Faculty Department Transfer Workflow & Governance Rules", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("When a faculty member is transferred between departments (e.g. from CSE to AI & DS), the system executes an automated, multi-layered synchronization process:")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    transfer_steps = [
        ("1. Transfer History Audit Log (staff_department_history): ", "An immutable record is inserted storing staff_id, from_dept, to_dept, transfer_date, and system timestamp."),
        ("2. Academic Profile & Credentials Update: ", "Department column in staff_academics is updated to the new target department. If assigned as a Dept Admin (admin_dep), the admin scope updates automatically. Session JWT claims update upon next login."),
        ("3. Server Storage Folder Relocation (moveFacultyDirectory): ", "Uploaded files on server disk are moved from /SREC/{old_dept}/{staff_id}/ to /SREC/{new_dept}/{staff_id}/. Database document paths across all activity tables are remapped automatically."),
        ("4. Historical Activity & Appraisal Retention: ", "All prior Publications, Patents, Grants, Certifications, and submitted FPI Appraisal forms remain preserved under staff_id without data loss."),
        ("5. HOD Review Authority Transfer: ", "The faculty member is removed from the old HOD's active directory and pending queue, and immediately appears in the new HOD's directory and pending appraisal queue."),
        ("6. Dynamic HOD Resolution in Reports: ", "General Information tables, FPI forms, and Dossier Reports dynamically resolve and display the HOD Name of the faculty's new department.")
    ]

    for title_text, body_text in transfer_steps:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(title_text)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(3, 105, 161)
        r2 = bp.add_run(body_text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(14)

    # Section 4
    add_heading_styled(doc, "4. Annual Performance Appraisal (FPI) Governance & Rules", level=1)

    fpi_rules = [
        ("Completed Academic Year Rule: ", "Performance appraisals evaluate the completed academic year (e.g. 2025-2026). The appraisal form defaults to getAppraisalAcademicYear() (2025-2026)."),
        ("Score Evaluation Caps (200 Total): ", "PART A (Teaching: Max 60 Pts), PART B (Prof Dev: Max 40 Pts), PART C (R&D: Max 80 Pts), PART D (Institutional: Max 20 Pts). Total FPI score capped at 200 Pts."),
        ("FPI Part D Scoring Formula & Worked Examples: ", "Part D evaluates Institutional and Departmental contributions: Institutional Roles award 10.0 Pts per duty (Max 20.0 Pts); Departmental Roles award 10.0 Pts per duty (Max 10.0 Pts). Total Part D score is computed as: min(20.0, Institutional_Marks + Departmental_Marks). Worked Example A: Faculty serving as Club Coordinator (10 inst) + Dept Exam Cell Coordinator (10 dept) -> 10 + 10 = 20.0 / 20.0 (Max Cap Reached). Worked Example B: Faculty holding 2 Institutional Roles (NSS Coordinator + Sports Advisory, 2 x 10 = 20 inst, 0 dept) -> 20.0 / 20.0 (Max Cap Reached)."),
        ("Club Coordinator & Co-Coordinator Rule: ", "Being assigned as a Club Coordinator or Club Co-Coordinator is categorized as an Institutional Level Additional Responsibility under Criteria D1 in PART D (Institutional Development & Contribution), scoring 10 Pts per duty up to the 20 Pts Part D cap. Furthermore, student club activities organized log extension points under PART B."),
        ("Designation-Based Customization: ", "System Admins / HR can define custom unit marks, max caps, calculation rules, and bracket configs per designation (Assistant Professor, Associate Professor, Professor, Professor & Head). System falls back to ALL common default mappings if no override exists."),
        ("Threshold Bracket Configurator (⚙️ Config Bracket): ", "Admins can configure cutoff thresholds (e.g. feedback 4.0 cutoff, pass % 80% cutoff, journal vs conference split, patent status splits)."),
        ("Custom PART Addition (➕ Add New PART): ", "Admins can dynamically add new evaluation sections (PART_E, PART_F) and modify section titles in real time."),
        ("Approval Workflow Lifecycle: ", "Submitted (Pending HOD Review) -> HOD Approved (Pending Principal/HR Review) -> Final Approved (Finalized Appraisal)."),
        ("Proof Document Verification Provision: ", "HODs, Principal, HR, and System Admin have access to the Auto-Mapped Activity Verification Panel inside both the submissions list and View Full FPI Form modal. Each record features an 👁️ View Proof button to inspect original evidence documents (PDF, images) before approving scores."),
        ("Bulk Appraisal Digital Signing & Batch Approval (HOD & System Admin): ", "HODs and System Admins (along with Principal & HR) can select multiple pending appraisal forms using multi-select checkboxes or 'Select All Pending' to review and execute bulk signing and approvals. The system applies cryptographic digital signatures (signer name, ISO timestamp, client IP address), sets evaluated marks matching rubrics/self-scores, logs audit timestamps (hod_approved_at / final_approved_at), and dispatches automated email notifications with score breakdowns to all approved faculty.")
    ]

    for title_text, body_text in fpi_rules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(title_text)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(15, 51, 31)
        r2 = bp.add_run(body_text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(14)

    # Section 5
    add_heading_styled(doc, "5. AI Academic Bio-Data & Real-Time Web Push Notification Governance", level=1)

    push_and_cv_rules = [
        ("1-Click AI CV & Statutory Bio-Data Generation: ", "Provides automated profile aggregation across 14 database tables. Supports 3 statutory formats: (1) SREC Official Letterhead Format with college crest and QR verification, (2) AICTE / Anna University Inspection Format with teaching/research splits, and (3) Modern Technical Europass Format. Features an intelligent AI Academic Statement Generator with dynamic tone selection (Executive, Research, Teaching) and section visibility toggles."),
        ("PWA Web Push Notification Infrastructure (VAPID): ", "Implements zero-cost, real-time push alerts via Service Worker (sw-push.js) and VAPID cryptography. Delivers instant alerts directly to faculty mobile devices and desktop browsers for appraisal review approvals, bulk digital signing, revision requests, and administrative circulars even when the browser tab is closed."),
        ("Profile Picture Background White Standardization Rule: ", "When a faculty member uploads a profile picture (JPG/PNG), the system automatically processes the image via AI portrait segmentation (rembg) to isolate the portrait and standardize the background to solid pure white (#FFFFFF) by default. This ensures uniform, professional presentation across all portals, ID cards, and official CV templates.")
    ]

    for title_text, body_text in push_and_cv_rules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(title_text)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(3, 105, 161)
        r2 = bp.add_run(body_text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(14)

    # Section 6: SREC FIS V3.1 Features
    add_heading_styled(doc, "6. SREC FIS V3.1 Workflow Enhancements Governance & Security", level=1)

    v31_rules = [
        ("AI Batch Document Upload & Concurrent Pre-Fill Pipeline (V3.1-01): ", "Faculty can upload up to 10 academic documents (Max 5 MB each, Max 20 MB total) across supported MIME types (PDF, JPG, PNG, WEBP). Employs a 2-worker queue concurrency to balance OCR and API rate limits. Partial failure resiliency ensures successful extractions are immediately reviewable while failed items offer retry/manual entry. In strict accordance with AI Non-Autonomy rules, zero unconfirmed records are auto-saved to database."),
        ("One-Click Consolidated Department Academic & Accreditation PDF Compilation (V3.1-02): ", "Authorized HODs and Administrators can generate a consolidated 9-section official Department Performance & Accreditation PDF in under 3.5 seconds. Strictly enforces department data isolation: HODs can only access their assigned department; cross-department attempts trigger HTTP 403 Forbidden. System Admin & Principal retain institutional scope.")
    ]

    for title_text, body_text in v31_rules:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(title_text)
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(14)
        r1.font.color.rgb = RGBColor(15, 51, 31)
        r2 = bp.add_run(body_text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(14)

    # Save document
    docs_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../docs'))
    os.makedirs(docs_dir, exist_ok=True)
    out_path = os.path.join(docs_dir, "System_Constraints_and_Portal_Rules.docx")
    doc.save(out_path)
    print(f"System constraints Word document successfully generated at: {out_path}")

if __name__ == "__main__":
    create_system_constraints_document()
