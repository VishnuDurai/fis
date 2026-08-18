import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def add_page_border(section, color="0F331F", sz="8"):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for b_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '20')
        b.set(qn('w:color'), color)
        pgBorders.append(b)
    sectPr.append(pgBorders)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(3, 105, 161)
    return p

def add_diagram_box(doc, title, flow_steps):
    add_heading_styled(doc, f"❖ Technical Architecture Diagram: {title}", level=2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r_t = p.add_run(f"CODEBASE ARCHITECTURE DIAGRAM: {title.upper()}\n")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(14)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(15, 51, 31)
    
    for idx, step in enumerate(flow_steps):
        p_step = cell.add_paragraph()
        p_step.paragraph_format.space_before = Pt(2)
        p_step.paragraph_format.space_after = Pt(4)
        
        step_prefix = f"► Component {idx+1}: "
        r_sp = p_step.add_run(step_prefix)
        r_sp.bold = True
        r_sp.font.name = 'Times New Roman'
        r_sp.font.size = Pt(14)
        r_sp.font.color.rgb = RGBColor(3, 105, 161)
        
        r_sb = p_step.add_run(step)
        r_sb.font.name = 'Times New Roman'
        r_sb.font.size = Pt(14)
        r_sb.font.color.rgb = RGBColor(30, 41, 59)
        
        if idx < len(flow_steps) - 1:
            p_arrow = cell.add_paragraph()
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_arrow.paragraph_format.space_before = Pt(0)
            p_arrow.paragraph_format.space_after = Pt(0)
            r_arr = p_arrow.add_run("│\n▼")
            r_arr.bold = True
            r_arr.font.name = 'Times New Roman'
            r_arr.font.size = Pt(14)
            r_arr.font.color.rgb = RGBColor(15, 118, 110)

    set_table_borders(tbl, color="0F331F", sz="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def create_technical_modification_guide():
    doc = Document()
    
    # Margins & Page Border
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        add_page_border(section, color="0F331F", sz="8")

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(6)
        footer_run = footer_p.add_run("© 2026 FIS Team - Sri Ramakrishna Engineering College, Coimbatore | SREC FIS V3.0")
        footer_run.font.name = 'Times New Roman'
        footer_run.font.size = Pt(11)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)

    # Official SREC Header Banner (same as PDF Reports)
    header_banner = os.path.abspath(os.path.join(os.path.dirname(__file__), '../client/public/srec-header-banner.png'))
    if os.path.exists(header_banner):
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_banner.paragraph_format.space_before = Pt(0)
        p_banner.paragraph_format.space_after = Pt(6)
        r_banner = p_banner.add_run()
        r_banner.add_picture(header_banner, width=Inches(7.0))

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sys.paragraph_format.space_before = Pt(0)
    p_sys.paragraph_format.space_after = Pt(10)
    r_sys = p_sys.add_run("FACULTY INFORMATION SYSTEM (SREC FIS V3.0)")
    r_sys.font.name = 'Times New Roman'
    r_sys.font.size = Pt(14)
    r_sys.bold = True
    r_sys.font.color.rgb = RGBColor(2, 132, 199)

    # Document Main Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("TECHNICAL CONSTRAINTS & CODEBASE FILE MODIFICATION GUIDE")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(14)
    run_meta = meta_p.add_run(f"Document Generated: {datetime.datetime.now().strftime('%B %d, %Y')} | Version 3.0 | Status: Active Developer Guide")
    run_meta.font.name = 'Times New Roman'
    run_meta.font.size = Pt(11)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Section 1
    add_heading_styled(doc, "1. System Overview & Developer Quick Reference", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("This technical manual specifies the exact code files, functions, and line ranges required to modify portal rules, menu visibility, role permissions, appraisal scoring algorithms, faculty transfer mechanisms, and file upload limits in the SREC FIS codebase.")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    add_diagram_box(doc, "Four-Tier Codebase Architecture & Data Flow", [
        "Frontend Tier (React 19 + Vite): Renders dynamic UI, Sidebar menu tree, SearchableSelect inputs, and Appraisal Form Engine.",
        "API Gateway & Routing Tier (Express.js): Authenticates JWT Bearer tokens and executes role permission guards.",
        "Database Storage Tier (MySQL 9.7+ / srec_fis): Stores 35+ relational tables, indexes, and historical audit logs.",
        "Physical Disk Directory Tier (/server/SREC/): Stores uploaded proof documents organized under dedicated /SREC/{Dept}/{Staff_ID}/ directories."
    ])

    # Table 1: Feature to File Mapping
    table1 = doc.add_table(rows=1, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.autofit = False

    t1_headers = ["Constraint / Rule Feature", "Target File Path", "Key Function / Block", "Line Range"]
    t1_widths = [Inches(2.2), Inches(2.4), Inches(1.8), Inches(0.8)]

    for i, h in enumerate(t1_headers):
        table1.rows[0].cells[i].text = h
        table1.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table1.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table1.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table1.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(14)
        set_cell_background(table1.rows[0].cells[i], "0F331F")
        set_cell_margins(table1.rows[0].cells[i])
        table1.rows[0].cells[i].width = t1_widths[i]

    mappings = [
        ("Navigation Sidebar Items & Tree", "client/src/components/Sidebar.jsx", "getMenuStructure()", "L60-L325"),
        ("Navbar Role Labels & Switcher", "client/src/components/Navbar.jsx", "roleLabel / user menu", "L120-L135"),
        ("Role JWT Issuance & Claims", "server/routes/auth.js", "POST /api/auth/login", "L35-L130"),
        ("File Access JWT Guard", "server/server.js", "requireFileAuth middleware", "L28-L44"),
        ("Faculty Transfer & Folder Move", "server/routes/admin.js", "POST /api/admin/faculty/transfer", "L715-L755"),
        ("Physical Disk Directory Relocation", "server/utils/fileStorage.js", "moveFacultyDirectory()", "L230-L290"),
        ("Department Transfer History Schema", "server/db.js", "staff_department_history DDL", "L554-L561"),
        ("Appraisal Completed Academic Year", "client/src/utils/academicYear.js", "getAppraisalAcademicYear()", "L20-L33"),
        ("Appraisal Engine & Designation Overrides", "client/src/pages/Appraisal.jsx", "getConstraint()", "L3565-L3610"),
        ("Designation Filter & Copy Overrides", "client/src/pages/Appraisal.jsx", "Target Designation Bar", "L1585-L1640"),
        ("Threshold Bracket Configurator", "client/src/pages/Appraisal.jsx", "ruleModalItem modal", "L5120-L5210"),
        ("Custom PART / Section Addition", "client/src/pages/Appraisal.jsx", "showAddPartModal & distinctSections", "L5215-L5310"),
        ("Appraisal Approval Workflow & Remarks", "client/src/pages/Appraisal.jsx", "viewingAppraisal submit", "L3470-L3550"),
        ("Bulk Appraisal Sign & Approve UI & Modal", "client/src/pages/Appraisal.jsx", "handleBulkSignAndApproveSubmit", "L210-L290, L4900-L5500"),
        ("Bulk HOD Appraisal Sign & Approve API", "server/routes/faculty.js", "POST /appraisals/bulk-hod-sign-approve", "L1730-L1840"),
        ("Bulk Final Appraisal Sign & Approve API", "server/routes/faculty.js", "POST /appraisals/bulk-final-sign-approve", "L1845-L1950"),
        ("1-Click AI CV & Bio-Data Generator UI", "client/src/pages/CVGenerator.jsx", "CVGenerator component & templates", "L1-L650"),
        ("Faculty CV Aggregator & AI Bio API", "server/routes/faculty.js", "GET /api/faculty/cv-data/:staffId", "L3460-L3600"),
        ("PWA Web Push Notification Manager", "client/src/components/NotificationPrompt.jsx", "NotificationPrompt component", "L1-L320"),
        ("Web Push VAPID & Dispatch API", "server/routes/notifications.js", "sendPushNotification()", "L1-L165"),
        ("Auto-Mapped Proof Document Links", "client/src/pages/Appraisal.jsx", "AutoMappedVerificationPanel", "L1227-L1490"),
        ("Faculty Read-Only Inputs (Dept Admin)", "client/src/pages/Personal.jsx", "disabled={auth.role === 'dept_admin'}", "L505-L545"),
        ("Responsibilities Hierarchy", "client/src/pages/Responsibilities.jsx", "isInstitutionalAdmin / isHod", "L25-L40"),
        ("Database Schema Definitions", "server/db.js", "tables DDL array", "L70-L640"),
        ("Profile Picture Upload & White BG API", "server/routes/activities.js", "POST /upload/profile-pic", "L372-L395"),
        ("AI Portrait Background Processor", "server/utils/processProfilePic.js", "standardizeProfilePic()", "L1-L50"),
        ("Python Portrait Isolation Script", "scripts/process_profile_picture.py", "process_image() (rembg + Pillow)", "L1-L85"),
        ("Auto-Doc Schema Script", "scripts/generate_schema_doc.py", "generate_schema_document()", "L1-L430"),
        ("Auto-Doc Constraints Script", "scripts/generate_system_constraints_doc.py", "create_system_constraints_doc()", "L1-L240"),
        ("Auto-Doc Technical Guide Script", "scripts/generate_tech_file_guide_doc.py", "create_technical_guide()", "L1-L250"),
        ("Auto-Doc Workflows Script", "scripts/generate_portal_workflows_doc.py", "create_portal_workflows_document()", "L1-L320")
    ]

    for row_idx, (feat, fpath, ffunc, flines) in enumerate(mappings):
        row_cells = table1.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate([feat, fpath, ffunc, flines]):
            row_cells[i].text = val
            p_cell = row_cells[i].paragraphs[0]
            p_cell.runs[0].font.name = 'Times New Roman'
            p_cell.runs[0].font.size = Pt(14)
            if i == 1 or i == 3:
                p_cell.runs[0].font.bold = True
                p_cell.runs[0].font.color.rgb = RGBColor(3, 105, 161)
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i])
            row_cells[i].width = t1_widths[i]

    set_table_borders(table1)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Save document
    docs_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../docs'))
    os.makedirs(docs_dir, exist_ok=True)
    out_path = os.path.join(docs_dir, "Technical_Constraints_and_File_Modification_Guide.docx")
    doc.save(out_path)
    print(f"Technical modification guide Word document successfully generated at: {out_path}")

if __name__ == "__main__":
    create_technical_modification_guide()
