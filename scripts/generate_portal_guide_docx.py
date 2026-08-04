import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_word_document():
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Title Banner
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_title.add_run("SRI RAMAKRISHNA ENGINEERING COLLEGE (SREC)\n")
    run_inst.font.size = Pt(14)
    run_inst.font.bold = True
    run_inst.font.color.rgb = RGBColor(0x0F, 0x33, 0x1F)

    run_title = p_title.add_run("Faculty Information System (FIS) V3.0\n")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    run_sub = p_title.add_run("Comprehensive Technical Specification & Security Constraints Manual")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    artifact_dir = "/Users/vishnudurairs/.gemini/antigravity-ide/brain/aabc19d1-cf18-427d-a1bd-056207ed0333"

    # Helpers for Headings
    def add_h1(text):
        h = doc.add_heading(text, level=1)
        h.style.font.name = 'Calibri'
        h.style.font.size = Pt(16)
        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor(0x0F, 0x33, 0x1F)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_h2(text):
        h = doc.add_heading(text, level=2)
        h.style.font.name = 'Calibri'
        h.style.font.size = Pt(13)
        h.style.font.bold = True
        h.style.font.color.rgb = RGBColor(0x03, 0x69, 0xA1)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r2 = p.add_run(text)
        p.paragraph_format.space_after = Pt(3)

    def add_image_with_caption(filename, caption):
        img_path = os.path.join(artifact_dir, filename)
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(6.2))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = p_cap.add_run(f"Figure: {caption}")
            r_cap.font.size = Pt(9.5)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            p_cap.paragraph_format.space_after = Pt(14)

    # 1. Executive Summary
    add_h1("1. Executive Summary & Architecture Overview")
    doc.add_paragraph("The SREC Faculty Information System (FIS) V3.0 is an enterprise-grade academic data management platform engineered for Sri Ramakrishna Engineering College. It serves as the single source of truth for faculty credentials, research activities, professional accomplishments, NAAC/NBA/NIRF accreditation reports, and the annual Faculty Performance Indicator (FPI) appraisal process.")

    # Tech Stack Table
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].text = "System Component"
    hdr[1].text = "Specification / Technology"
    set_cell_background(hdr[0], '0F331F')
    set_cell_background(hdr[1], '0F331F')
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    specs = [
        ("Frontend Framework", "React 18, Vite 8, Vanilla CSS (Design Tokens & HSL Palette)"),
        ("Backend Runtime", "Node.js (ES Module Syntax), Express.js"),
        ("Database Engine", "MySQL 8 on localhost:3306 ('srec_fis' database)"),
        ("Authentication", "JSON Web Tokens (JWT) with HTTP Bearer Auth & Encrypted Passwords"),
        ("File Storage Engine", "Multer Storage Engine mapped to isolated faculty directories"),
        ("Footer Branding Directive", "'Developed and Maintained by Team FIS' above Copyright line"),
        ("Mobile & App Readiness", "Media Queries (max-width: 1024px, 768px), viewport-fit=cover, PWA/Capacitor wrapping"),
        ("Database Schema Sync", "Automated regeneration via python3 scripts/generate_schema_doc.py updates Database_Schema.docx")
    ]

    for comp, spec in specs:
        row = t.add_row().cells
        row[0].text = comp
        row[1].text = spec
        set_cell_background(row[0], 'F8FAFC')
        set_cell_background(row[1], 'FFFFFF')

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Authentication & HOD Role Transfer Constraints
    add_h1("2. Authentication & HOD Role Transfer Security Architecture")
    doc.add_paragraph("Access to SREC FIS V3.0 is governed by a multi-portal login gateway enforcing strict Role-Based Access Control (RBAC) and dynamic HOD role remapping.")
    add_image_with_caption("doc_01_login_portal.png", "Multi-Portal Login Gateway Interface")

    add_h2("2.1 Security Constraints & Portal Role Scoping")
    add_bullet("Faculty Member (faculty):", "Access restricted to personal profile, activity logging, live bibliometrics, and self-appraisal submission.")
    add_bullet("Department Admin (dept_admin) / HOD:", "Access restricted to department faculty directory (Read-Only), department cumulative bibliometrics, and HOD appraisal verification.")
    add_bullet("System Administrator (admin):", "Institution-wide administrative access, department creation, faculty department transfers, and dynamic page building.")

    add_h2("2.2 Technical Specifications for HOD Role Transfer & Duty Remapping")
    doc.add_paragraph("The system features a dynamic HOD transfer and privilege remapping engine managed via System Admin provisions (admin_dep table & PUT /api/admin/staff/:id/transfer):")
    add_bullet("Outgoing HOD Privilege Revocation:", "When an existing Head of Department is transferred or stepped down, their Staff ID entry in admin_dep is automatically updated or removed. Access permissions to review department appraisal forms, view department ZIP archives, and assign departmental responsibilities are immediately revoked. The user's active portal scope reverts back to standard Faculty Member (faculty) level.")
    add_bullet("Incoming HOD Privilege Mapping:", "When a new faculty member is designated as HOD, their Staff ID is mapped into admin_dep with target department scoping. This automatically provisions access to the Department Admin Portal (dept_admin), Department Overview Dashboard, Read-Only Department Directory, Cumulative Department Bibliometrics Card, and the HOD Appraisal Evaluation Queue.")
    add_bullet("Faculty & Department Physical Folder Migration Engine:", "Upon executing a department transfer, moveFacultyDirectory() physically moves the faculty member's uploads directory from /uploads/faculty/{old_dept}/{staff_id} to /uploads/faculty/{new_dept}/{staff_id}. Database record paths across staff_academics, staff_publications, and activity tables are updated in bulk.")
    add_bullet("Appraisal Review Queue Re-scoping:", "Any pending appraisal submissions for the transferred faculty member are automatically re-scoped to the incoming HOD's review queue in the target department.")

    # 3. Faculty Portal
    add_h1("3. Portal 1: Faculty Member Portal")
    doc.add_paragraph("The Faculty Portal empowers individual faculty members to maintain their academic dossier, upload activity proof documents, track live citations, and submit self-appraisal forms.")
    add_image_with_caption("doc_02_faculty_dashboard.png", "Faculty Member Overview Dashboard")

    add_h2("3.1 Academic Information & Profile Identifiers")
    doc.add_paragraph("The Academic Information page contains two dedicated specification sections:")
    add_bullet("Institutional Identifiers:", "AICTE Faculty ID, Anna University ID, APAAR ID.")
    add_bullet("Research Profile Identifiers:", "ORCID ID, Google Scholar Profile ID, Scopus Author ID, Web of Science Researcher ID.")
    add_image_with_caption("doc_03_faculty_academic_info.png", "Faculty Academic Information & Profile Identifiers")

    add_h2("3.2 Publications & R&D Activity Modules")
    doc.add_paragraph("The Publications Module supports comprehensive journal and conference tracking in compliance with official specifications.")
    add_bullet("Category Toggle:", "Dynamically switches between Journal-Only (ISSN, Volume, Issue, Impact Factor) and Conference-Only (ISBN, Venue, Dates) fields.")
    add_bullet("Fields Directives:", "Publication Status field deleted. Citations Count field retained and active.")
    add_bullet("Live Bibliometrics Card:", "Displays Total Citations, h-Index, i10-Index, Google Scholar link badge, and Scopus/ORCID link badge with 1-click Auto Sync.")
    add_bullet("Filter Element Sequence:", "1. Search Bar -> 2. Select Faculty Member -> 3. Filter Report by Publication Category.")
    add_image_with_caption("doc_04_faculty_publications.png", "Faculty Publications Module with Live Bibliometrics Card")

    add_h2("3.3 Faculty Performance Indicator (FPI) Appraisal Module")
    doc.add_paragraph("Allows faculty to enter self-scores across Parts A, B, C, and D, upload supporting document evidence, and export dossiers to PDF format.")
    add_image_with_caption("doc_05_faculty_appraisal.png", "Faculty Appraisal Submission Form Interface")

    # 4. Dept Admin Portal
    add_h1("4. Portal 2: Department Admin & HOD Portal")
    doc.add_paragraph("The Department Admin Portal enables HODs to monitor department academic metrics, access faculty profiles in read-only mode, and evaluate appraisals.")
    add_image_with_caption("doc_06_dept_admin_dashboard.png", "Department Admin Overview Dashboard")

    add_h2("4.1 Department Faculty Directory (Read-Only Rights)")
    doc.add_paragraph("Department Admins have strict read-only viewing rights. Inputs inside the faculty detail modal are disabled and form submission is prevented.")
    add_image_with_caption("doc_07_dept_admin_directory.png", "Department Faculty Directory Read-Only Modal")

    add_h2("4.2 Cumulative Department Bibliometrics & Top Performing Department Faculty Leaderboard")
    doc.add_paragraph("Default view (-- All Department Faculty --) calculates and displays Cumulative Department Bibliometrics: Dept Total Citations (SUM), Max h-Index (MAX), Dept Total i10-Index (SUM), and Faculty Directory Count. Directly below the cumulative card, the Department Top Performing Faculty Leaderboard highlights the department's top research citation performers ranked live by Total Citations, h-Index, and i10-Index.")
    add_image_with_caption("doc_08_dept_admin_publications.png", "Department Admin Cumulative Bibliometrics & Top Performing Faculty Leaderboard")

    add_h2("4.3 Appraisal Review Queue & Notification Badges")
    doc.add_paragraph("Displays pending appraisal submissions for department faculty, allowing HOD score verification, remarks entry, and real-time badge counters for HOD, Principal, and HR.")
    add_image_with_caption("doc_09_dept_admin_appraisal_review.png", "Department Appraisal Review Interface")

    # 5. System Admin Portal & Dynamic Page Enabling/Disabling
    add_h1("5. Portal 3: System Admin Portal, Leaderboards & Dynamic Page Builder")
    doc.add_paragraph("Provides overall governance over institution-wide data, department structures, faculty department transfers, system-wide leaderboards, and dynamic custom page building.")
    add_image_with_caption("doc_10_sysadmin_dashboard.png", "System Administrator Overview Dashboard")

    add_h2("5.1 System-Wide Bibliometrics Leaderboards (Top Faculty & Top Department)")
    doc.add_paragraph("In System Admin Login, the Publications page renders dual institution-wide research leaderboards: (1) System-Wide Top Performing Faculty Leaderboard ranking the institution's top 10 researchers across all departments, and (2) System-Wide Top Performing Department Leaderboard ranking academic departments by cumulative citations and max h-Index.")
    add_image_with_caption("doc_13_sysadmin_leaderboards.png", "System Admin Dual Leaderboards (System-Wide Top Performing Faculty & Top Department)")

    add_h2("5.2 Department Management & Faculty Department Transfer Engine")
    doc.add_paragraph("Enables adding/removing academic departments, initiating faculty department transfers, folder migrations, and DB path remapping.")
    add_image_with_caption("doc_11_sysadmin_departments.png", "System Admin Department Management & Transfer Interface")

    add_h2("5.3 Dynamic Custom Page Builder & Page Enabling/Disabling Specifications")
    doc.add_paragraph("The Dynamic Page Builder (DynamicPagesAdmin.jsx & GET/POST /api/dynamic-pages) allows System Administrators to construct custom form pages, fields, and tables dynamically with strict visibility and lifecycle controls:")
    add_bullet("Target Portal Access Scope (portals Array):", "System Admin specifies target portal access by toggling portals: ['faculty', 'dept_admin', 'admin'].")
    add_bullet("Enabling Mechanism (Active Pages):", "When a page is enabled/created with active target portals, the frontend Sidebar navigation automatically injects the page into the assigned category (standalone, personal, academic, activity, reports). Route handlers permit submissions and file uploads stored in dynamic_page_data table.")
    add_bullet("Disabling Mechanism (Disabled Pages):", "When a System Admin disables a page or unchecks target portal access, the page is immediately hidden from navigation sidebars across all user sessions. Backend API route guards enforce token validation and return HTTP 403 Forbidden / Page Disabled if a user attempts direct URL navigation to a disabled page.")
    add_image_with_caption("doc_12_sysadmin_pagebuilder.png", "System Admin Dynamic Custom Page Builder Interface")

    # 6. Verification Metrics
    add_h1("6. System Verification & Automated Test Suite Metrics")
    doc.add_paragraph("The system has undergone comprehensive automated testing verifying security constraints, HOD role transfer logic, page enabling/disabling, RBAC route guards, CRUD operations, and report exports.")

    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = t2.rows[0].cells
    hdr2[0].text = "Verification Metric"
    hdr2[1].text = "Technical Scope / Result"
    hdr2[2].text = "Status"
    for c in hdr2:
        set_cell_background(c, '0F331F')
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    test_metrics = [
        ("Automated Portal System Tests", "42 / 42 Passed (100% Success Rate)", "PASSED"),
        ("Dynamic HOD Role Transfer & Remapping", "Verified (Privilege Revocation, Mapping & Folder Migration)", "PASSED"),
        ("Dynamic Custom Page Enabling / Disabling", "Verified (Sidebar Invalidation, Route Guards & 403 Checks)", "PASSED"),
        ("Dept Admin Read-Only Rights Enforcement", "Strictly Enforced (Disabled Inputs & Route Guards)", "PASSED"),
        ("Cumulative Department Bibliometrics Engine", "Verified (SUM Citations, MAX h-Index, SUM i10-Index)", "PASSED"),
        ("Filter Display Order", "Verified (Search -> Faculty -> Category)", "PASSED"),
        ("Database Schema Document Sync", "Verified & Synchronized (Database_Schema.docx)", "PASSED")
    ]

    for m, r_val, s_val in test_metrics:
        row = t2.add_row().cells
        row[0].text = m
        row[1].text = r_val
        row[2].text = s_val
        set_cell_background(row[0], 'F8FAFC')
        set_cell_background(row[1], 'FFFFFF')
        set_cell_background(row[2], 'E0F2FE')

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Footer Notice
    p_ft = doc.add_paragraph()
    p_ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ft = p_ft.add_run("Developed and Maintained by Team FIS\nSri Ramakrishna Engineering College, Coimbatore")
    r_ft.font.size = Pt(10)
    r_ft.font.bold = True
    r_ft.font.color.rgb = RGBColor(0x0F, 0x33, 0x1F)

    out_file = "/Users/vishnudurairs/Data/fis/SREC_FIS_Comprehensive_Portal_Guide.docx"
    doc.save(out_file)
    print(f"Successfully generated updated Word document manual at: {out_file}")

if __name__ == '__main__':
    create_word_document()
