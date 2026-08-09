import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches

# Schema section mappings & table descriptions
SECTION_MAPPINGS = [
    {
        "category": "1. Authentication & System Administration",
        "description": "Tables managing administrator credentials, department admin accounts, and faculty login credentials.",
        "tables": ["admin", "admin_dep", "staff_user"]
    },
    {
        "category": "2. Faculty Personal & Academic Profiles",
        "description": "Core identity profiles, contact info, AICTE/Anna Univ IDs, experience tracking, qualifications, and document files.",
        "tables": ["staff_personal", "staff_academics", "staff_edu", "staff_member", "staff_pan", "staff_aadhar"]
    },
    {
        "category": "3. Research, Publications & Intellectual Property",
        "description": "Records of journal/conference publications, books, patents, grants, consultancy revenue, and Ph.D. guidance.",
        "tables": ["staff_publication", "staff_book_published", "staff_ipr", "staff_funding", "staff_development", "staff_scholars", "staff_supervisor", "staff_seed_money"]
    },
    {
        "category": "4. Events, Certifications & Professional Activities",
        "description": "FDPs/workshops attended or organized, resource person roles, club events, certifications, exams, and awards.",
        "tables": ["staff_interaction", "staff_resource", "staff_event_organized", "staff_club", "staff_award", "staff_certificate", "staff_competitive", "staff_innovative"]
    },
    {
        "category": "5. Performance Appraisals & Department Governance",
        "description": "Faculty self-appraisals (annual evaluations), FPI criteria templates, department level duty responsibilities, and transfer history logs.",
        "tables": ["staff_appraisal", "appraisal_template", "staff_responsibilities", "staff_department_history"]
    },
    {
        "category": "6. Master Lookup Tables",
        "description": "Standardized drop-down and lookup lists for departments, universities, professional bodies, designations, and clubs.",
        "tables": ["departments", "university", "professional", "designations", "clubs"]
    }
]

TABLE_DESCRIPTIONS = {
    "admin": "Stores system administrator credentials.",
    "admin_dep": "Stores department-level administrator credentials and department scope.",
    "staff_user": "Stores faculty login credentials and profile picture references.",
    "staff_personal": "Stores personal, contact, identification numbers (PAN, Aadhaar, AICTE, Anna Univ, APAAR), and identity documents.",
    "staff_academics": "Stores academic designation, department, qualification, and detailed breakdown of past and SREC experience.",
    "staff_edu": "Stores educational qualifications (UG, PG, Ph.D.), degrees, institutes, marks, and certificate file references.",
    "staff_member": "Stores memberships in professional societies (IEEE, CSI, ACM, etc.).",
    "staff_pan": "Upload path references for PAN card documents.",
    "staff_aadhar": "Upload path references for Aadhaar card documents.",
    "staff_publication": "Stores research papers published in journals and conferences with indexing and citation metrics.",
    "staff_book_published": "Stores books and book chapters published by faculty.",
    "staff_ipr": "Stores patents filed, published, or granted with institutional affiliations.",
    "staff_funding": "Stores research grants, funding agency details, sanctioned amounts, and project statuses.",
    "staff_development": "Stores consultancy projects and research & development revenue generation.",
    "staff_scholars": "Stores details of Ph.D. scholars guided by faculty members.",
    "staff_supervisor": "Stores Ph.D. supervisor recognition status and scholar counts.",
    "staff_interaction": "Stores FDPs, workshops, and seminars attended by faculty.",
    "staff_resource": "Stores roles acted as resource person, session chair, or guest speaker.",
    "staff_event_organized": "Stores seminars, workshops, and conferences organized by faculty.",
    "staff_club": "Stores club events and student chapter activities coordinated.",
    "staff_award": "Stores honors, awards, and recognitions received.",
    "staff_certificate": "Stores online courses completed (NPTEL, Coursera, etc.) and scores.",
    "staff_competitive": "Stores competitive exam achievements (GATE, NET, SET, etc.).",
    "staff_innovative": "Stores innovative teaching methodologies and project initiatives.",
    "staff_appraisal": "Stores annual faculty self-appraisal submissions and evaluation metrics.",
    "appraisal_template": "Stores master criteria, rubrics, and maximum marks for faculty performance index (FPI).",
    "staff_responsibilities": "Stores institutional and departmental responsibilities assigned to faculty.",
    "staff_department_history": "Stores faculty department transfer history log including effective transfer dates.",
    "staff_seed_money": "Stores internal seed money grants awarded for research projects.",
    "departments": "Lookup table for institutional academic departments and official acronyms.",
    "university": "Lookup table for standard university names.",
    "professional": "Lookup table for recognized professional societies.",
    "designations": "Lookup table for academic designations.",
    "clubs": "Lookup table for campus student clubs and technical societies with Faculty Incharge & Co-Faculty Incharge assignments."
}

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    border = OxmlElement('w:insideV')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
    tblPr.append(tblBorders)

def add_page_border(section, color="0F331F", sz="8"):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for b_name in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '20')
        b.set(qn('w:color'), color)
        pgBorders.append(b)
    sectPr.append(pgBorders)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(3, 105, 161)
    return p

def generate_er_diagram_image(img_path):
    """Dynamically renders a high-resolution Entity-Relationship (ER) Diagram image."""
    fig, ax = plt.subplots(figsize=(16, 12), dpi=300)
    ax.set_xlim(0, 160)
    ax.set_ylim(0, 120)
    ax.axis('off')
    fig.patch.set_facecolor('#FFFFFF')

    # Main Diagram Header Title
    ax.text(80, 115, "SREC FIS V3.0 — ENTITY-RELATIONSHIP (ER) ARCHITECTURE DIAGRAM", 
            fontsize=15, fontweight='bold', color='#0F331F', ha='center', fontfamily='DejaVu Sans')
    ax.text(80, 111, "Relational Entity Mapping Across 35+ MySQL Database Tables & Key Foreign Keys", 
            fontsize=11, fontstyle='italic', color='#475569', ha='center', fontfamily='DejaVu Sans')

    # Helper box drawer
    def draw_entity_box(x, y, w, h, title, fields, color_hdr='#0F331F', color_bg='#F8FAFC'):
        # Outer box
        rect = patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.3,rounding_size=0.8", 
                                      facecolor=color_bg, edgecolor='#CBD5E1', linewidth=1.2)
        ax.add_patch(rect)
        
        # Header bar
        hdr_rect = patches.FancyBboxPatch((x, y + h - 4.5), w, 4.5, boxstyle="round,pad=0.3,rounding_size=0.4", 
                                          facecolor=color_hdr, edgecolor='none')
        ax.add_patch(hdr_rect)
        
        # Header text
        ax.text(x + w/2, y + h - 2.5, title, fontsize=9.5, fontweight='bold', color='#FFFFFF', 
                ha='center', va='center', fontfamily='DejaVu Sans')
        
        # Fields list
        curr_y = y + h - 7.5
        for f in fields:
            ax.text(x + 1.5, curr_y, f, fontsize=8, color='#1E293B', fontfamily='DejaVu Sans')
            curr_y -= 3.0

    # Cluster 1: Core Identity & Credentials (Top Left)
    draw_entity_box(5, 75, 32, 30, "staff_personal [PK: staff_id]", [
        "[PK] staff_id (VARCHAR)",
        "• staff_name, gender, dob",
        "• email, mobile, address",
        "• pan_no, aadhar_no",
        "• aicte_id, anna_univ_id",
        "• is_relieved (TINYINT)"
    ], color_hdr='#0F331F', color_bg='#F0FDF4')

    draw_entity_box(42, 75, 30, 30, "staff_academics", [
        "[FK] staff_id (-> staff_personal)",
        "• Department, Designation",
        "• Qualification, Specialization",
        "• doj_srec, exp_srec_yrs",
        "• past_exp_yrs, total_exp",
        "• orcid_id, scholar_id, scopus"
    ], color_hdr='#0F331F', color_bg='#F0FDF4')

    draw_entity_box(5, 48, 32, 22, "staff_user", [
        "[FK] staff_id (-> staff_personal)",
        "• username, password (bcrypt)",
        "• profile_pic, created_at"
    ], color_hdr='#166534', color_bg='#F0FDF4')

    draw_entity_box(42, 48, 30, 22, "admin & admin_dep", [
        "[PK] id, [FK] staff_id",
        "• username, password",
        "• Department (admin_dep)",
        "• role: admin / dept_admin"
    ], color_hdr='#166534', color_bg='#F0FDF4')

    # Cluster 2: Performance Appraisal & Governance (Top Right)
    draw_entity_box(80, 75, 34, 30, "staff_appraisal [FPI Engine]", [
        "[PK] id, [FK] staff_id",
        "• academic_year (e.g. 2025-2026)",
        "• part_a_score, part_b_score",
        "• part_c_score, part_d_score",
        "• total_fpi_score (Max 200)",
        "• status: Submitted / Approved"
    ], color_hdr='#991B1B', color_bg='#FEF2F2')

    draw_entity_box(120, 75, 34, 30, "appraisal_template", [
        "[PK] id, criteria_code",
        "• section_code (PART_A..D)",
        "• criteria_title, max_marks",
        "• unit_mark, target_designation",
        "• bracket_config (JSON)"
    ], color_hdr='#991B1B', color_bg='#FEF2F2')

    draw_entity_box(80, 48, 34, 22, "staff_responsibilities", [
        "[PK] id, [FK] staff_id",
        "• responsibility, level (Inst/Dept)",
        "• assigned_by, department",
        "• academic_year"
    ], color_hdr='#B91C1C', color_bg='#FEF2F2')

    draw_entity_box(120, 48, 34, 22, "staff_department_history", [
        "[PK] id, [FK] staff_id",
        "• from_dept, to_dept",
        "• transfer_date, timestamp"
    ], color_hdr='#B91C1C', color_bg='#FEF2F2')

    # Cluster 3: Research & IP (Middle Row Left)
    draw_entity_box(5, 12, 34, 30, "staff_publication & books", [
        "[PK] id, [FK] staff_id",
        "• title, journal_name, issn",
        "• pub_type: Journal / Conf",
        "• indexing: Scopus / WoS / SCI",
        "• book_title, isbn, publisher"
    ], color_hdr='#0369A1', color_bg='#F0F9FF')

    draw_entity_box(43, 12, 34, 30, "staff_ipr & funding", [
        "[PK] id, [FK] staff_id",
        "• patent_title, status, date",
        "• project_title, agency",
        "• sanctioned_amount, granted",
        "• staff_seed_money (title, amt)"
    ], color_hdr='#0369A1', color_bg='#F0F9FF')

    # Cluster 4: Activities & Clubs (Middle Row Right)
    draw_entity_box(81, 12, 35, 30, "clubs & staff_club", [
        "[PK] id",
        "• name (Club Name)",
        "• [FK] faculty_incharge_id",
        "• [FK] co_faculty_incharge_id",
        "• staff_club (events organized)"
    ], color_hdr='#0F766E', color_bg='#F0FDFA')

    draw_entity_box(120, 12, 34, 30, "staff_interaction & certs", [
        "[PK] id, [FK] staff_id",
        "• FDP / workshop_title, duration",
        "• staff_resource (talk_title)",
        "• staff_certificate (NPTEL, etc)",
        "• staff_award (award_name)"
    ], color_hdr='#0F766E', color_bg='#F0FDFA')

    # Cluster Connectors / Cardinality Lines
    # staff_personal to staff_academics
    ax.annotate('', xy=(42, 90), xytext=(37, 90),
                arrowprops=dict(arrowstyle="->,head_width=0.3", color="#0F331F", lw=1.5))
    ax.text(39.5, 91.5, "1:1", fontsize=8.5, fontweight='bold', color="#0F331F", fontfamily='DejaVu Sans')

    # staff_personal to staff_appraisal
    ax.annotate('', xy=(80, 90), xytext=(72, 90),
                arrowprops=dict(arrowstyle="->,head_width=0.3", color="#991B1B", lw=1.5))
    ax.text(76, 91.5, "1:N", fontsize=8.5, fontweight='bold', color="#991B1B", fontfamily='DejaVu Sans')

    # staff_personal to staff_publication
    ax.annotate('', xy=(20, 42), xytext=(20, 48),
                arrowprops=dict(arrowstyle="->,head_width=0.3", color="#0369A1", lw=1.5))
    ax.text(21, 45, "1:N", fontsize=8.5, fontweight='bold', color="#0369A1", fontfamily='DejaVu Sans')

    # staff_personal to clubs
    ax.annotate('', xy=(81, 27), xytext=(77, 27),
                arrowprops=dict(arrowstyle="->,head_width=0.3", color="#0F766E", lw=1.5))
    ax.text(78.5, 28.5, "1:N", fontsize=8.5, fontweight='bold', color="#0F766E", fontfamily='DejaVu Sans')

    # Save figure image
    plt.tight_layout()
    plt.savefig(img_path, format='png', dpi=300, bbox_inches='tight')
    plt.close(fig)

def generate_docx():
    docs_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../docs'))
    os.makedirs(docs_dir, exist_ok=True)
    output_path = os.path.join(docs_dir, 'Database_Schema.docx')
    er_img_path = os.path.join(docs_dir, 'er_diagram.png')
    
    # Render fresh ER Diagram image
    try:
        generate_er_diagram_image(er_img_path)
        print(f"ER Diagram image generated successfully at: {er_img_path}")
    except Exception as e:
        print(f"Warning: Failed to generate ER diagram image: {e}")

    use_mysql = False
    tables_data = {}
    
    try:
        import pymysql
        conn = pymysql.connect(
            host='localhost',
            port=3306,
            user='root',
            password='',
            database='srec_fis',
            cursorclass=pymysql.cursors.DictCursor
        )
        cursor = conn.cursor()
        cursor.execute("SHOW TABLES")
        tables_res = cursor.fetchall()
        tbl_key = list(tables_res[0].keys())[0] if tables_res else 'Tables_in_srec_fis'
        existing_tables = set(r[tbl_key] for r in tables_res)
        
        for tname in existing_tables:
            cursor.execute(f"DESCRIBE `{tname}`")
            cols = cursor.fetchall()
            formatted_cols = []
            for cid, col in enumerate(cols):
                is_pk = col['Key'] == 'PRI'
                formatted_cols.append((
                    cid,
                    col['Field'],
                    col['Type'],
                    col['Null'] == 'NO',
                    col['Default'],
                    1 if is_pk else 0
                ))
            tables_data[tname] = formatted_cols
            
        use_mysql = True
        conn.close()
        print("Connected to MySQL 'srec_fis' database successfully.")
    except Exception as e:
        print(f"MySQL Connection fallback: {e}")
        import sqlite3
        db_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '../server/fis.db'))
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%';")
        existing_tables = set(row[0] for row in cursor.fetchall())
        for tname in existing_tables:
            cursor.execute(f"PRAGMA table_info('{tname}');")
            tables_data[tname] = cursor.fetchall()
        conn.close()

    doc = Document()
    
    # Page setup - Margins & Page Border
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        add_page_border(section, color="0F331F", sz="8")
        
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(6)
        footer_run = footer_p.add_run("© 2026 FIS Team - Sri Ramakrishna Engineering College, Coimbatore | SREC FIS V3.0")
        footer_run.font.name = 'Times New Roman'
        footer_run.font.size = Pt(11)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)
        
    # Official SREC Header Banner (same as PDF Reports)
    header_banner = os.path.abspath(os.path.join(os.path.dirname(__file__), '../client/public/srec-header-banner.png'))
    if os.path.exists(header_banner):
        p_banner = doc.add_paragraph()
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_banner.paragraph_format.space_before = Pt(0)
        p_banner.paragraph_format.space_after = Pt(6)
        r_banner = p_banner.add_run()
        r_banner.add_picture(header_banner, width=Inches(7.0))

    p_sys = doc.add_paragraph()
    p_sys.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sys.paragraph_format.space_before = Pt(0)
    p_sys.paragraph_format.space_after = Pt(10)
    r_sys = p_sys.add_run("FACULTY INFORMATION SYSTEM (SREC FIS V3.0)")
    r_sys.font.name = 'Times New Roman'
    r_sys.font.size = Pt(14)
    r_sys.bold = True
    r_sys.font.color.rgb = RGBColor(2, 132, 199)

    # Document Main Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("DATABASE ARCHITECTURE & TABLE SCHEMA REFERENCE")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(16)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    # Metadata Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    now_str = datetime.datetime.now().strftime("%B %d, %Y - %H:%M:%S")
    engine_name = "MySQL 9.7+ (srec_fis)" if use_mysql else "SQLite 3 (fis.db)"
    
    meta_data = [
        [("Database Engine:", True), (engine_name, False)],
        [("Last Synchronized:", True), (now_str, False)]
    ]
    
    for row_idx, row in enumerate(meta_data):
        for col_idx, (text, is_bold) in enumerate(row):
            cell = meta_table.cell(row_idx, col_idx)
            cell.width = Inches(3.5)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = is_bold
            run.font.color.rgb = RGBColor(30, 41, 59)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            
    set_table_borders(meta_table)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Insert High-Resolution Entity-Relationship (ER) Diagram Image
    if os.path.exists(er_img_path):
        add_heading_styled(doc, "1. Entity-Relationship (ER) Architecture Diagram", level=1)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(10)
        r_img = p_img.add_run()
        r_img.add_picture(er_img_path, width=Inches(6.8))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(14)
        r_cap = p_cap.add_run("Figure 1: High-Resolution Entity-Relationship (ER) Diagram for SREC FIS V3.0 Relational Database")
        r_cap.font.name = 'Times New Roman'
        r_cap.font.size = Pt(11)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(71, 85, 105)

    processed_tables = set()
    existing_tables = set(tables_data.keys())
    
    # Iterate through structured sections
    for sec_idx, section_info in enumerate(SECTION_MAPPINGS):
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(4)
        h1.paragraph_format.keep_with_next = True
        run_h1 = h1.add_run(section_info["category"])
        run_h1.font.name = 'Times New Roman'
        run_h1.font.size = Pt(16)
        run_h1.font.bold = True
        run_h1.font.color.rgb = RGBColor(15, 23, 42)
        
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(10)
        run_desc = p_desc.add_run(section_info["description"])
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(14)
        run_desc.font.italic = True
        run_desc.font.color.rgb = RGBColor(71, 85, 105)
        
        for tname in section_info["tables"]:
            if tname in tables_data:
                processed_tables.add(tname)
                cols_info = tables_data[tname]
                
                h2 = doc.add_paragraph()
                h2.paragraph_format.space_before = Pt(12)
                h2.paragraph_format.space_after = Pt(2)
                h2.paragraph_format.keep_with_next = True
                
                r_name = h2.add_run(f"Table: {tname}")
                r_name.font.name = 'Times New Roman'
                r_name.font.size = Pt(14)
                r_name.font.bold = True
                r_name.font.color.rgb = RGBColor(3, 105, 161)
                
                if tname in TABLE_DESCRIPTIONS:
                    r_tdesc = h2.add_run(f" — {TABLE_DESCRIPTIONS[tname]}")
                    r_tdesc.font.name = 'Times New Roman'
                    r_tdesc.font.size = Pt(14)
                    r_tdesc.font.color.rgb = RGBColor(71, 85, 105)
                
                # Render table
                tbl = doc.add_table(rows=1, cols=6)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                
                hdr_cells = tbl.rows[0].cells
                headers = ["#", "Column Name", "Data Type", "Nullable", "Default Value", "Key"]
                col_widths = [Inches(0.4), Inches(2.2), Inches(1.5), Inches(0.9), Inches(1.3), Inches(0.7)]
                
                for i, head_text in enumerate(headers):
                    hdr_cells[i].text = head_text
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(14)
                    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    set_cell_background(hdr_cells[i], "0F331F")
                    set_cell_margins(hdr_cells[i])
                    hdr_cells[i].width = col_widths[i]
                
                for row_idx, col in enumerate(cols_info):
                    row_cells = tbl.add_row().cells
                    cid, cname, ctype, notnull, dflt, pk = col[0], col[1], col[2], col[3], col[4], col[5]
                    
                    is_not_null = notnull if isinstance(notnull, bool) else (notnull == 1)
                    null_str = "No" if is_not_null else "Yes"
                    key_str = "PK" if pk else ""
                    dflt_str = str(dflt) if dflt is not None else "NULL"
                    
                    vals = [str(cid+1), cname, str(ctype).upper(), null_str, dflt_str, key_str]
                    bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
                    
                    for c_i, v in enumerate(vals):
                        row_cells[c_i].text = v
                        p = row_cells[c_i].paragraphs[0]
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run()
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                        if c_i == 1 and pk:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(185, 28, 28)
                        else:
                            run.font.color.rgb = RGBColor(30, 41, 59)
                        set_cell_background(row_cells[c_i], bg_color)
                        set_cell_margins(row_cells[c_i])
                        row_cells[c_i].width = col_widths[c_i]
                
                set_table_borders(tbl)
                doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.save(output_path)
    print(f"Database schema Word document generated successfully at: {output_path}")

if __name__ == '__main__':
    generate_docx()
